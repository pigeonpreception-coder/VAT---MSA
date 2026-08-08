from __future__ import annotations

import csv
import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "VAT-MSA_Enterprise_Architecture_Blueprint.docx"


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def audit_docx() -> dict[str, object]:
    require(DOCX.exists(), "DOCX output is missing")
    require(zipfile.is_zipfile(DOCX), "DOCX is not a valid ZIP package")
    with zipfile.ZipFile(DOCX) as archive:
        bad = archive.testzip()
        require(bad is None, f"Corrupt DOCX member: {bad}")

    document = Document(DOCX)
    require(len(document.sections) == 1, "Expected one document section")
    section = document.sections[0]
    require(section.page_width == 7_772_400 and section.page_height == 10_058_400, "Page size is not US Letter")
    require(
        all(value == 914_400 for value in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin)),
        "Margins are not one inch",
    )

    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    for marker in ("TODO", "TBD", "FIXME", "[[PAGEBREAK]]", "CONTINUE"):
        require(marker not in all_text, f"Placeholder marker remains in DOCX: {marker}")

    require(len(document.tables) == 24, f"Expected 24 body tables, found {len(document.tables)}")
    for index, table in enumerate(document.tables, start=1):
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        require(tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9_360, f"Table {index} width is not 9360 DXA")
        grid = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        require(sum(grid) == 9_360, f"Table {index} grid sums to {sum(grid)} DXA")
        for row_number, row in enumerate(table.rows, start=1):
            require(len(row.cells) == len(grid), f"Table {index}, row {row_number} has inconsistent cell count")
            cell_widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                require(tc_w is not None, f"Table {index}, row {row_number} has missing cell width")
                cell_widths.append(int(tc_w.get(qn("w:w"))))
            require(cell_widths == grid, f"Table {index}, row {row_number} widths do not match table grid")

    images = []
    for shape in document.inline_shapes:
        doc_pr = shape._inline.docPr
        description = doc_pr.get("descr") or ""
        title = doc_pr.get("title") or ""
        require(description.strip(), "Inline image is missing alt-text description")
        require(title.strip(), "Inline image is missing alt-text title")
        images.append({"title": title, "description": description})
    require(len(images) == 8, f"Expected 8 architecture figures, found {len(images)}")

    headings = [p.text for p in document.paragraphs if p.style and p.style.name.startswith("Heading")]
    require(any(text.startswith("18. Governance") for text in headings), "Governance section heading is missing")
    require(any(text.startswith("Appendix H.") for text in headings), "Source appendix heading is missing")

    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(images),
        "headings": len(headings),
        "sha256": sha256(DOCX),
        "bytes": DOCX.stat().st_size,
    }


def audit_artifacts() -> dict[str, object]:
    json_schema = ROOT / "03-api" / "schemas" / "vat-msa-invoice.schema.json"
    schema = json.loads(json_schema.read_text(encoding="utf-8"))
    require(schema.get("$schema") == "https://json-schema.org/draft/2020-12/schema", "Unexpected JSON Schema dialect")
    require(schema.get("type") == "object", "Invoice schema root is not an object")

    openapi = (ROOT / "03-api" / "openapi.yaml").read_text(encoding="utf-8")
    require(re.search(r"^openapi:\s*3\.1", openapi, flags=re.MULTILINE) is not None, "OpenAPI document is not version 3.1")
    require("vat-msa-invoice.schema.json" in openapi, "OpenAPI document does not reference the invoice schema")
    require("\t" not in openapi, "OpenAPI YAML contains tab indentation")

    events = (ROOT / "03-api" / "event-catalog.yaml").read_text(encoding="utf-8")
    require("cloudevents" in events.lower(), "Event catalogue does not declare CloudEvents")
    require("\t" not in events, "Event catalogue YAML contains tab indentation")

    schema_sql = (ROOT / "04-data" / "core-schema.sql").read_text(encoding="utf-8")
    table_count = len(re.findall(r"^CREATE TABLE\s+", schema_sql, flags=re.IGNORECASE | re.MULTILINE))
    require(table_count >= 20, f"Expected at least 20 reference tables, found {table_count}")
    for invariant in ("vat_ledger_entry", "outbox_event", "idempotency_record", "audit_event"):
        require(invariant in schema_sql, f"Core SQL schema is missing {invariant}")

    with (ROOT / "05-security" / "rbac-matrix.csv").open(encoding="utf-8-sig", newline="") as stream:
        rows = list(csv.DictReader(stream))
    require(len(rows) >= 20, f"RBAC matrix is unexpectedly small: {len(rows)} rows")
    require(
        {"role", "resource", "action", "scope", "approval_required", "segregation_of_duties_note"}.issubset(rows[0]),
        "RBAC matrix columns are incomplete",
    )

    diagram_dir = ROOT / "02-diagrams"
    require(len(list(diagram_dir.glob("*.mmd"))) == 8, "Expected 8 Mermaid source diagrams")
    require(len(list(diagram_dir.glob("*.png"))) == 8, "Expected 8 rendered diagrams")

    required = [
        ROOT / "README.md",
        ROOT / "04-data" / "data-dictionary.md",
        ROOT / "05-security" / "security-controls-matrix.md",
        ROOT / "06-delivery" / "non-functional-requirements.md",
        ROOT / "06-delivery" / "testing-strategy.md",
        ROOT / "06-delivery" / "roadmap.md",
        ROOT / "references" / "architecture-decisions.md",
        ROOT / "references" / "assumptions-and-open-decisions.md",
    ]
    missing = [str(path.relative_to(ROOT)) for path in required if not path.exists()]
    require(not missing, f"Required package files are missing: {missing}")

    return {
        "reference_tables": table_count,
        "rbac_rows": len(rows),
        "mermaid_diagrams": 8,
        "rendered_diagrams": 8,
    }


def main() -> int:
    result = {"docx": audit_docx(), "artifacts": audit_artifacts()}
    print(json.dumps(result, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"AUDIT_FAILED: {exc}", file=sys.stderr)
        raise
