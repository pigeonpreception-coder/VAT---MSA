from __future__ import annotations

import math
import re
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
SOURCE_MD = BASE / "01-blueprint" / "VAT-MSA_Enterprise_Architecture_Blueprint.md"
DIAGRAM_DIR = BASE / "02-diagrams"
OUTPUT = BASE / "VAT-MSA_Enterprise_Architecture_Blueprint.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "1B7F8C"
INK = "1F2937"
GRAY = "5E6B78"
MID_GRAY = "A8B2BD"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
GOLD = "B98522"
RED = "9B1C1C"
GREEN = "246B4B"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="C9D2DC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths sum to {sum(widths)}, expected {CONTENT_WIDTH_DXA}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.width = Inches(widths[idx] / 1440)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text_node, end])
    set_run_font(run, size=8.5, color=GRAY)


def set_paragraph_shading_and_border(paragraph, fill=PALE_BLUE, border_color=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(sz)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://[^\s)]+)")


def add_markdown_runs(paragraph, text: str, base_size=11, base_color=INK):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=max(8.5, base_size - 1), color=DARK_BLUE)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            run._r.get_or_add_rPr().append(shd)
        else:
            add_hyperlink(paragraph, token, token)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    # Reuse Word's schema-valid built-in List Bullet and List Number abstract
    # definitions, then patch their level-0 geometry to the selected preset.
    bullet_num_id = int(doc.styles["List Bullet"]._element.xpath("./w:pPr/w:numPr/w:numId/@w:val")[0])
    decimal_num_id = int(doc.styles["List Number"]._element.xpath("./w:pPr/w:numPr/w:numId/@w:val")[0])

    def abstract_for_num(num_id: int) -> int:
        num = next(x for x in numbering.findall(qn("w:num")) if int(x.get(qn("w:numId"))) == num_id)
        return int(num.find(qn("w:abstractNumId")).get(qn("w:val")))

    def patch_level(abstract_id: int):
        abstract = next(x for x in numbering.findall(qn("w:abstractNum")) if int(x.get(qn("w:abstractNumId"))) == abstract_id)
        lvl = next(x for x in abstract.findall(qn("w:lvl")) if int(x.get(qn("w:ilvl"))) == 0)
        p_pr = lvl.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            lvl.append(p_pr)
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for child in list(tabs):
            tabs.remove(child)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")

    bullet_abstract_id = abstract_for_num(bullet_num_id)
    decimal_abstract_id = abstract_for_num(decimal_num_id)
    patch_level(bullet_abstract_id)
    patch_level(decimal_abstract_id)
    return bullet_num_id, decimal_abstract_id


def create_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def distribute_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    cols = len(rows[0])
    scores = []
    for idx in range(cols):
        lengths = [len(re.sub(r"\*\*|`", "", row[idx])) for row in rows]
        mx = max(lengths)
        avg = sum(lengths) / max(1, len(lengths))
        score = min(48, max(8, avg * 0.65 + mx * 0.35))
        if idx == 0 and cols >= 3:
            score *= 0.85
        scores.append(score)
    min_width = 720 if cols >= 6 else 900
    available = CONTENT_WIDTH_DXA - min_width * cols
    total_score = sum(scores)
    widths = [min_width + int(available * score / total_score) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: Sequence[Sequence[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = distribute_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_markdown_runs(p, text, base_size=9.0, base_color=NAVY if r_idx == 0 else INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.widow_control = True
    add_markdown_runs(p, text)
    return p


def add_list_paragraph(doc: Document, text: str, num_id: int, level=0):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, num_id, level)
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    add_markdown_runs(p, text)
    return p


def add_callout(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.13)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = True
    set_paragraph_shading_and_border(p)
    add_markdown_runs(p, text, base_size=10.4, base_color=NAVY)
    return p


def add_code_block(doc: Document, lines: Sequence[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading_and_border(p, fill="F7F8FA", border_color="CBD5E1")
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Courier New", size=8.0, color=INK)
        if idx < len(lines) - 1:
            run.add_break()
    return p


def set_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc: Document, filename: str, caption: str):
    path = DIAGRAM_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.25))
    set_alt_text(shape, caption.split(".")[0], caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=GRAY, italic=True)


def populate_running_header(header):
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6100, 3260], indent_dxa=0)
    remove_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(left.add_run("VAT-MSA Enterprise Architecture Blueprint"), size=8.5, color=GRAY, bold=True)
    set_run_font(right.add_run("Version 1.0 | 8 August 2026"), size=8.5, color=GRAY)


def populate_running_footer(footer):
    footer.is_linked_to_previous = False
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [7000, 2360], indent_dxa=0)
    remove_table_borders(ftable)
    for cell in ftable.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    fl = ftable.cell(0, 0).paragraphs[0]
    fr = ftable.cell(0, 1).paragraphs[0]
    fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fl.add_run("PROPOSED - FOR NAMRA VALIDATION"), size=8.0, color=GRAY, bold=True)
    set_run_font(fr.add_run("Page "), size=8.5, color=GRAY)
    add_page_field(fr)


def add_header_footer(section):
    section.different_first_page_header_footer = True
    populate_running_header(section.header)

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].text = ""

    populate_running_footer(section.footer)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(fp.add_run("Architecture baseline | 8 August 2026"), size=8.5, color=GRAY)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = rgb(GRAY)
    caption.font.italic = True

    code = styles.add_style("VAT-MSA Code", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8)
    code.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("NATIONAL TAX TECHNOLOGY PLATFORM"), size=10.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("VAT-MSA"), size=34, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("Enterprise Architecture Blueprint"), size=23, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run_font(p.add_run("VAT Transaction, Reconciliation, Compliance and Audit Platform"), size=13.5, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.65)
    p.paragraph_format.right_indent = Inches(0.65)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(32)
    set_paragraph_shading_and_border(p, fill=PALE_BLUE, border_color=BLUE)
    set_run_font(p.add_run("A governed implementation baseline for discovery, legal validation, procurement and national-scale delivery."), size=11.5, color=NAVY, bold=True)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2350, 7010])
    set_table_borders(meta, color="D5DDE6", size="4")
    labels = ["Status", "Version", "Prepared", "Audience"]
    values = [
        "Proposed for validation and Architecture Review Board approval",
        "1.0 architecture baseline",
        "8 August 2026",
        "NamRA, government stakeholders and delivery partners",
    ]
    for idx, (label, value) in enumerate(zip(labels, values)):
        set_cell_shading(meta.cell(idx, 0), LIGHT_GRAY)
        for cell in meta.rows[idx].cells:
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = meta.cell(idx, 0).paragraphs[0]
        vp = meta.cell(idx, 1).paragraphs[0]
        set_run_font(lp.add_run(label), size=9.5, color=NAVY, bold=True)
        set_run_font(vp.add_run(value), size=9.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    set_run_font(p.add_run("PROPOSED - FOR NAMRA VALIDATION"), size=10, color=RED, bold=True)
    doc.add_page_break()


def parse_markdown(doc: Document, md: str, bullet_num_id: int, decimal_abstract_id: int):
    # The cover is rendered separately. Start at the first substantive control section.
    start = md.find("## Document control")
    if start == -1:
        raise ValueError("Document control marker not found")
    lines = md[start:].splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    active_decimal_num_id = None
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            idx += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            idx += 1
            continue

        if stripped == "[[PAGEBREAK]]":
            active_decimal_num_id = None
            doc.add_page_break()
            idx += 1
            continue

        fig = re.fullmatch(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", stripped)
        if fig:
            active_decimal_num_id = None
            add_figure(doc, fig.group(1), fig.group(2))
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[idx + 1]):
            active_decimal_num_id = None
            table_lines = [stripped]
            idx += 2  # skip separator
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = []
            for raw in table_lines:
                cells = [c.strip() for c in raw.strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            continue

        if stripped.startswith("### "):
            active_decimal_num_id = None
            p = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(p, stripped[4:], base_size=12, base_color=DARK_BLUE)
            idx += 1
            continue
        if stripped.startswith("## "):
            active_decimal_num_id = None
            p = doc.add_paragraph(style="Heading 2")
            add_markdown_runs(p, stripped[3:], base_size=13, base_color=BLUE)
            idx += 1
            continue
        if stripped.startswith("# "):
            active_decimal_num_id = None
            p = doc.add_paragraph(style="Heading 1")
            add_markdown_runs(p, stripped[2:], base_size=16, base_color=BLUE)
            idx += 1
            continue

        if stripped.startswith("> "):
            active_decimal_num_id = None
            add_callout(doc, stripped[2:])
            idx += 1
            continue

        bullet = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet:
            active_decimal_num_id = None
            level = min(len(bullet.group(1)) // 2, 2)
            add_list_paragraph(doc, bullet.group(2), bullet_num_id, level)
            idx += 1
            continue

        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered:
            level = min(len(numbered.group(1)) // 2, 2)
            if active_decimal_num_id is None:
                active_decimal_num_id = create_num_instance(doc, decimal_abstract_id)
            add_list_paragraph(doc, numbered.group(2), active_decimal_num_id, level)
            idx += 1
            continue

        active_decimal_num_id = None
        add_body_paragraph(doc, stripped.rstrip("  "))
        idx += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)


def pil_font(size: int, bold=False):
    candidates = []
    if bold:
        candidates.extend([Path("C:/Windows/Fonts/calibrib.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")])
    else:
        candidates.extend([Path("C:/Windows/Fonts/calibri.ttf"), Path("C:/Windows/Fonts/arial.ttf")])
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


class Diagram:
    def __init__(self, title: str, width=1800, height=980):
        self.width = width
        self.height = height
        self.image = Image.new("RGB", (width, height), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.draw.rounded_rectangle((28, 28, width - 28, height - 28), radius=24, outline="#D6DEE7", width=3, fill="#FFFFFF")
        self.draw.text((70, 58), title, font=pil_font(36, True), fill="#0B2545")
        self.draw.line((70, 112, width - 70, 112), fill="#2E74B5", width=4)

    def text_wrap(self, text: str, font, max_width: int) -> list[str]:
        out = []
        for para in text.split("\n"):
            words = para.split()
            if not words:
                out.append("")
                continue
            line = words[0]
            for word in words[1:]:
                trial = line + " " + word
                if self.draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                    line = trial
                else:
                    out.append(line)
                    line = word
            out.append(line)
        return out

    def box(self, xy, text, fill="#E8EEF5", outline="#2E74B5", font_size=25, bold=False, radius=18, align="center"):
        x1, y1, x2, y2 = xy
        self.draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
        font = pil_font(font_size, bold)
        lines = self.text_wrap(text, font, max(40, x2 - x1 - 34))
        line_gap = int(font_size * 0.28)
        heights = [self.draw.textbbox((0, 0), line or " ", font=font)[3] for line in lines]
        total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
        y = y1 + (y2 - y1 - total_h) / 2
        for line, h in zip(lines, heights):
            bbox = self.draw.textbbox((0, 0), line, font=font)
            if align == "left":
                x = x1 + 18
            else:
                x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
            self.draw.text((x, y), line, font=font, fill="#1F2937")
            y += h + line_gap

    def label(self, xy, text, size=22, color="#5E6B78", bold=False, anchor="la"):
        self.draw.text(xy, text, font=pil_font(size, bold), fill=color, anchor=anchor)

    def arrow(self, start, end, color="#5E6B78", width=5, dashed=False):
        x1, y1 = start
        x2, y2 = end
        if dashed:
            steps = 14
            for i in range(0, steps, 2):
                a = i / steps
                b = min(1, (i + 1) / steps)
                self.draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill=color, width=width)
        else:
            self.draw.line((x1, y1, x2, y2), fill=color, width=width)
        angle = math.atan2(y2 - y1, x2 - x1)
        length = 18
        spread = 0.55
        p1 = (x2, y2)
        p2 = (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread))
        p3 = (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread))
        self.draw.polygon([p1, p2, p3], fill=color)

    def group(self, xy, title, fill="#F8FAFC", outline="#A8B2BD"):
        self.draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=3)
        x1, y1, _, _ = xy
        self.draw.text((x1 + 24, y1 + 18), title, font=pil_font(24, True), fill="#0B2545")

    def save(self, path: Path):
        self.image.save(path, format="PNG", dpi=(180, 180), optimize=True)


def build_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    d = Diagram("VAT-MSA system context")
    d.group((65, 155, 470, 900), "Taxpayer ecosystem")
    for y, text in [(235, "POS and retail SaaS"), (375, "ERP and accounting"), (515, "Seller / buyer portals"), (655, "Offline desktop")]:
        d.box((105, y, 430, y + 92), text, fill="#F4F7FB")
    d.group((650, 155, 1150, 900), "VAT-MSA core", fill="#F3F8FC", outline="#2E74B5")
    for y, text in [(235, "API and integration gateway"), (365, "Invoice + VAT transaction core"), (495, "Matching, exceptions and risk"), (625, "VAT ledger, evidence and analytics")]:
        d.box((700, y, 1100, y + 100), text, fill="#E8EEF5", bold=(y == 365))
    d.group((1330, 155, 1735, 900), "Government ecosystem")
    for y, text in [(235, "ITAS / taxpayer accounts"), (375, "Customs and imports"), (515, "Payments and refunds"), (655, "IAM / PKI / SOC")]:
        d.box((1370, y, 1695, y + 92), text, fill="#F4F7FB")
    for y in (281, 421, 561, 701):
        d.arrow((430, y), (700, y), color="#2E74B5")
        d.arrow((1370, y), (1100, y), color="#1B7F8C")
    d.save(DIAGRAM_DIR / "01-system-context.png")

    d = Diagram("Seven-layer logical platform")
    layers = [
        ("1", "Experience", "Taxpayer, NamRA, administration and public applications"),
        ("2", "Edge and integration", "API gateway, adapters, file exchange and verification"),
        ("3", "Domain services", "Taxpayer, invoice, VAT, return, audit and risk"),
        ("4", "Transaction and events", "Workflow, rules, messaging and transactional outbox"),
        ("5", "Data", "Operational stores, sub-ledger, evidence and search"),
        ("6", "Intelligence", "Matching, anomaly detection, warehouse and graph"),
    ]
    y = 155
    fills = ["#E8EEF5", "#EDF4FA", "#E5F1F2", "#EEF3F7", "#F1F4F7", "#F6F3EA"]
    for idx, (num, name, detail) in enumerate(layers):
        d.box((90, y, 1400, y + 100), f"{num}. {name}\n{detail}", fill=fills[idx], bold=True, align="left")
        if idx < len(layers) - 1:
            d.arrow((745, y + 100), (745, y + 125), color="#A8B2BD", width=4)
        y += 125
    d.box((1460, 155, 1710, 880), "7. Platform and security\n\nIAM\nPKI / HSM\nObservability\nCI/CD\nDR", fill="#0B2545", outline="#0B2545", font_size=24, bold=True)
    # Overpaint the dark-box text in white for readability.
    d.draw.rounded_rectangle((1460, 155, 1710, 880), radius=18, fill="#0B2545", outline="#0B2545", width=3)
    lines = ["7. Platform", "and security", "", "IAM", "PKI / HSM", "Observability", "CI/CD", "DR"]
    yy = 235
    for line in lines:
        bbox = d.draw.textbbox((0, 0), line, font=pil_font(24, True))
        d.draw.text((1585 - (bbox[2]-bbox[0])/2, yy), line, font=pil_font(24, True), fill="white")
        yy += 62
    for yy in (205, 330, 455, 580, 705, 830):
        d.arrow((1460, yy), (1400, yy), color="#B98522", width=4)
    d.save(DIAGRAM_DIR / "02-platform-layers.png")

    d = Diagram("Controlled fiscal document lifecycle")
    main = [(100, "RECEIVED"), (360, "VALIDATING"), (660, "ACCEPTED"), (920, "CERTIFIED"), (1190, "POSTED")]
    for x, text in main:
        d.box((x, 260, x + 220, 360), text, fill="#E8EEF5", bold=True)
    for i in range(len(main)-1):
        d.arrow((main[i][0]+220, 310), (main[i+1][0], 310), color="#2E74B5")
    d.box((1490, 200, 1690, 300), "MATCHED", fill="#E8F3ED", outline="#246B4B", bold=True)
    d.box((1490, 395, 1690, 495), "EXCEPTION", fill="#FFF5E8", outline="#B98522", bold=True)
    d.arrow((1410, 310), (1490, 250), color="#246B4B")
    d.arrow((1410, 310), (1490, 445), color="#B98522")
    d.box((360, 565, 600, 665), "REJECTED", fill="#FBECEC", outline="#9B1C1C", bold=True)
    d.arrow((470, 360), (480, 565), color="#9B1C1C")
    d.arrow((360, 615), (220, 360), color="#9B1C1C", dashed=True)
    d.label((635, 610), "Correct and resubmit", size=21, color="#9B1C1C")
    d.box((1185, 610, 1435, 710), "REVERSED", fill="#FBECEC", outline="#9B1C1C", bold=True)
    d.arrow((1300, 360), (1310, 610), color="#9B1C1C")
    d.label((870, 450), "Schema, identity, duplicate, arithmetic and VAT rule controls", size=23, color="#5E6B78", anchor="ma")
    d.save(DIAGRAM_DIR / "03-invoice-lifecycle.png")

    d = Diagram("Certification transaction and event flow")
    stages = [
        (80, "Seller\nsystem"), (320, "API\ngateway"), (560, "Invoice\nservice"), (800, "Rules\nservice"),
        (1040, "Transaction\nDB"), (1280, "Outbox /\nbroker"), (1520, "Matching, audit\nand analytics")
    ]
    for x, text in stages:
        d.box((x, 175, x + 200, 285), text, fill="#F4F7FB", bold=True)
        d.draw.line((x + 100, 285, x + 100, 840), fill="#D6DEE7", width=3)
    messages = [
        (180, 420, 320, "Submit + idempotency"),
        (420, 520, 560, "Authorised request"),
        (660, 800, 800, "Evaluate rules"),
        (900, 760, 1040, "Rule evidence"),
        (660, 650, 1040, "Atomic invoice + ledger + outbox"),
        (1140, 760, 1280, "Publish committed events"),
        (1380, 700, 1520, "Match / audit / analyse"),
    ]
    for x1, y, x2, label in messages:
        d.arrow((x1, y), (x2, y), color="#2E74B5" if x2 > x1 else "#1B7F8C", width=4)
        d.label(((x1+x2)/2, y-14), label, size=18, color="#5E6B78", anchor="ms")
    d.arrow((560, 780), (320, 780), color="#246B4B", width=5)
    d.arrow((320, 820), (180, 820), color="#246B4B", width=5)
    d.label((370, 752), "Signed certification receipt", size=19, color="#246B4B", anchor="ms")
    d.save(DIAGRAM_DIR / "04-transaction-processing.png")

    d = Diagram("Two-sided VAT reconciliation graph")
    d.box((720, 160, 1080, 270), "CERTIFIED INVOICE", fill="#E8EEF5", bold=True)
    d.box((160, 330, 480, 430), "SELLER TAXPAYER", fill="#F4F7FB", bold=True)
    d.box((1320, 330, 1640, 430), "BUYER TAXPAYER", fill="#F4F7FB", bold=True)
    d.box((160, 560, 480, 660), "OUTPUT VAT ENTRY", fill="#EDF4FA")
    d.box((1320, 560, 1640, 660), "INPUT VAT CANDIDATE", fill="#EDF4FA")
    d.box((720, 515, 1080, 635), "MATCHING DECISION", fill="#E5F1F2", outline="#1B7F8C", bold=True)
    d.box((570, 760, 880, 855), "MATCHED → PERIOD / RETURN", fill="#E8F3ED", outline="#246B4B", bold=True)
    d.box((980, 760, 1290, 855), "EXCEPTION WORKFLOW", fill="#FFF5E8", outline="#B98522", bold=True)
    d.arrow((720, 220), (480, 380), color="#2E74B5")
    d.arrow((1080, 220), (1320, 380), color="#2E74B5")
    d.arrow((320, 430), (320, 560), color="#2E74B5")
    d.arrow((1480, 430), (1480, 560), color="#2E74B5")
    d.arrow((480, 610), (720, 575), color="#1B7F8C")
    d.arrow((1320, 610), (1080, 575), color="#1B7F8C")
    d.arrow((820, 635), (725, 760), color="#246B4B")
    d.arrow((980, 635), (1135, 760), color="#B98522")
    d.save(DIAGRAM_DIR / "05-reconciliation-graph.png")

    d = Diagram("Zero-trust security zones")
    zones = [
        ((75, 170, 370, 860), "Internet and partners", "Untrusted\nTaxpayer systems\nPublic verification", "#F7F8FA", "#A8B2BD"),
        ((420, 170, 720, 860), "Protected edge", "DDoS / WAF\nAPI gateway\nRate and schema policy", "#EDF4FA", "#2E74B5"),
        ((770, 170, 1070, 860), "Application zone", "Domain workloads\nWorkflow / events\nWorkload identity", "#E5F1F2", "#1B7F8C"),
        ((1120, 170, 1420, 860), "Restricted data", "Operational DB\nVAT sub-ledger\nEvidence storage", "#F6F3EA", "#B98522"),
        ((1470, 170, 1725, 860), "Privileged", "PAM\nCI/CD\nHSM / KMS\nSOC", "#FBECEC", "#9B1C1C"),
    ]
    for xy, title, body, fill, outline in zones:
        d.group(xy, title, fill=fill, outline=outline)
        x1, y1, x2, y2 = xy
        d.box((x1 + 25, y1 + 105, x2 - 25, y2 - 55), body, fill=fill, outline=outline, font_size=24)
    for x in (370, 720, 1070, 1420):
        d.arrow((x, 515), (x+50, 515), color="#5E6B78", width=5)
    d.label((900, 920), "Every hop: authenticate identity → authorise resource and purpose → log outcome", size=25, color="#0B2545", bold=True, anchor="ma")
    d.save(DIAGRAM_DIR / "06-security-trust-zones.png")

    d = Diagram("High availability and disaster recovery")
    d.group((80, 170, 845, 855), "Primary production site", fill="#F3F8FC", outline="#2E74B5")
    d.group((955, 170, 1720, 855), "Secondary DR site", fill="#F8FAFC", outline="#A8B2BD")
    primary = [(220, 270, "Ingress"), (220, 410, "Application cells"), (145, 590, "PostgreSQL HA"), (390, 590, "Broker"), (635, 590, "Evidence store")]
    secondary = [(1095, 270, "Standby ingress"), (1095, 410, "Warm application cells"), (1020, 590, "Replicated DB"), (1265, 590, "Broker mirror"), (1510, 590, "Evidence replica")]
    for x, y, text in primary:
        d.box((x, y, x+200, y+88), text, fill="#E8EEF5", bold=(y == 410))
    for x, y, text in secondary:
        d.box((x, y, x+200, y+88), text, fill="#F2F4F7", outline="#A8B2BD", bold=(y == 410))
    d.arrow((320, 358), (320, 410), color="#2E74B5")
    for x in (245, 490, 735):
        d.arrow((320, 498), (x, 590), color="#2E74B5")
    d.arrow((1195, 358), (1195, 410), color="#5E6B78")
    for x in (1120, 1365, 1610):
        d.arrow((1195, 498), (x, 590), color="#5E6B78")
    for y, x1, x2, label in [(650, 345, 1020, "Encrypted continuous replication"), (700, 590, 1265, "Durable state replication"), (750, 835, 1510, "Immutable evidence replication")]:
        d.arrow((x1, y), (x2, y), color="#1B7F8C", dashed=True)
        d.label(((x1+x2)/2, y-12), label, size=18, color="#1B7F8C", anchor="ms")
    d.label((900, 920), "Traffic management fails over only after controlled health, data and business validation", size=23, color="#0B2545", bold=True, anchor="ma")
    d.save(DIAGRAM_DIR / "07-deployment-and-dr.png")

    d = Diagram("Outcome-gated delivery roadmap", width=2000, height=900)
    phases = [
        ("G0", "Mandate + legal design"), ("P1", "Foundation"), ("P2", "E-invoicing pilot"), ("P3", "VAT transaction + ledger"),
        ("P4", "Reconciliation + returns"), ("P5", "Audit + risk + refunds"), ("P6", "Offline + broad rollout"), ("P7", "National scale + analytics")
    ]
    x = 70
    y = 330
    width = 205
    gap = 35
    for idx, (code, name) in enumerate(phases):
        fill = "#E8EEF5" if idx < 4 else ("#E5F1F2" if idx < 7 else "#F6F3EA")
        outline = "#2E74B5" if idx < 4 else ("#1B7F8C" if idx < 7 else "#B98522")
        d.box((x, y, x+width, y+165), f"{code}\n{name}", fill=fill, outline=outline, font_size=23, bold=True)
        if idx < len(phases)-1:
            d.arrow((x+width, y+82), (x+width+gap, y+82), color="#5E6B78", width=4)
        x += width + gap
    d.label((1000, 590), "Each phase exits through legal, integrity, security, operational and recovery evidence", size=26, color="#0B2545", bold=True, anchor="ma")
    d.label((1000, 690), "MVP proves certification → production proves transaction/return control → national scale follows measured adoption", size=23, color="#5E6B78", anchor="ma")
    d.save(DIAGRAM_DIR / "08-delivery-roadmap.png")


def build_document():
    build_diagrams()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_header_footer(section)
    bullet_num_id, decimal_abstract_id = create_numbering(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE_MD.read_text(encoding="utf-8"), bullet_num_id, decimal_abstract_id)

    props = doc.core_properties
    props.title = "VAT-MSA Enterprise Architecture Blueprint"
    props.subject = "National VAT transaction, reconciliation, compliance and audit platform"
    props.author = "VAT-MSA Architecture Programme"
    props.last_modified_by = "VAT-MSA Architecture Programme"
    props.keywords = "VAT-MSA, NamRA, enterprise architecture, electronic invoicing, VAT, Namibia"
    props.comments = "Proposed architecture baseline for validation."

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_document()
